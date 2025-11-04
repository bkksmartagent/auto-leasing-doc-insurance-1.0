import streamlit as st
from docx import Document
from docx.shared import Pt, Inches
#from docx2pdf import convert
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
import tempfile
import os
import re
from datetime import datetime, timedelta
from babel.dates import format_date
from num2words import num2words
from PIL import Image
from io import BytesIO

def replace_placeholders(doc: Document, data: dict):
    for paragraph in doc.paragraphs:
        for key, value in data.items():
            placeholder = f"{{{{{key}}}}}"
            if placeholder in paragraph.text:
                for run in paragraph.runs:
                    if placeholder in run.text:
                        run.text = run.text.replace(placeholder, str(value))


def generate_contract(data, template_path, output_path):
    doc = Document(template_path)

    for section in doc.sections:
        replace_text_in_paragraphs(section.header.paragraphs, data)

    replace_text_in_paragraphs(doc.paragraphs, data)
    replace_text_in_tables(doc.tables, data)

    if data.get("photo1"):
        replace_image_placeholder_in_paragraphs(doc.paragraphs, "{{photo1}}", data.get("photo1"))

    if data.get("photo2"):
        replace_image_placeholder_in_paragraphs(doc.paragraphs, "{{photo2}}", data.get("photo2"))

    doc.save(output_path)

def generate_booking(data, template_path, output_path):
    doc = Document(template_path)
    for section in doc.sections:
        replace_text_in_paragraphs(section.header.paragraphs, data)
    replace_text_in_paragraphs(doc.paragraphs, data)
    replace_text_in_tables(doc.tables, data)
    doc.save(output_path)

def generate_furniture(template_path, output_path, data, furniture_list):
    doc = Document(template_path)
    for section in doc.sections:
        replace_text_in_paragraphs(section.header.paragraphs, data)
    replace_text_in_paragraphs(doc.paragraphs, data)
    replace_text_in_tables(doc.tables, data)
    replace_image_placeholder_in_paragraphs(doc.paragraphs, "{{photo1}}", data.get("photo1"))
    replace_image_placeholder_in_paragraphs(doc.paragraphs, "{{photo2}}", data.get("photo2"))

    for item in furniture_list:
        if item.get("image"):
            with tempfile.NamedTemporaryFile(delete=False, suffix=".jpg") as tmp:
                tmp.write(item["image"].read())
                tmp.flush()
                item["image_path"] = tmp.name
            item["image"].seek(0)

    insert_furniture_table(doc, furniture_list)

    doc.save(output_path)

def insert_furniture_table(doc, furniture_list):
    for table in doc.tables:
        if len(table.columns) == 3:
            table.allow_autofit = False
            table.columns[0].width = Inches(0.5)
            table.columns[1].width = Inches(2.5)
            table.columns[2].width = Inches(3.0)

            for idx, item in enumerate(furniture_list, start=1):
                row = table.add_row().cells
                row[0].text = str(idx)

                if "image_path" in item:
                    paragraph = row[1].paragraphs[0]
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    run = paragraph.add_run()
                    run.add_picture(item["image_path"], width=Inches(2.0))  # ปรับขนาดรูป

                row[2].text = item.get("remark", "")

            break

def bahttext(amount):
    amount = float(amount)
    integer_part = int(amount)
    decimal_part = int(round((amount - integer_part) * 100))
    baht_word = num2words(integer_part, lang='th')
    satang_word = num2words(decimal_part, lang='th') + 'สตางค์' if decimal_part > 0 else 'ถ้วน'
    return f"{baht_word}บาท{satang_word}"

def ordinal(n):
    if 11 <= (n % 100) <= 13:
        suffix = 'th'
    else:
        suffix = {1: 'st', 2: 'nd', 3: 'rd'}.get(n % 10, 'th')
    return f"{n}{suffix}"

def replace_text_in_paragraphs(paragraphs, data):
    for p in paragraphs:
        for key, val in data.items():
            placeholder = f"{{{{{key}}}}}"
            if placeholder in p.text:
                if isinstance(val, BytesIO):
                    continue
                full_text = p.text.replace(placeholder, str(val))
                runs = list(p.runs)
                for run in runs:
                    run._element.getparent().remove(run._element)
                run = p.add_run(full_text)
                font = run.font
                font.name = 'Cordia New'
                r = run._element
                r.rPr.rFonts.set(qn('w:eastAsia'), 'Cordia New')
                font.size = Pt(15)


def replace_text_in_tables(tables, data):
    for table in tables:
        for row in table.rows:
            for cell in row.cells:
                replace_text_in_paragraphs(cell.paragraphs, data)
                if cell.tables:
                    replace_text_in_tables(cell.tables, data)

def replace_image_placeholder_in_paragraphs(paragraphs, placeholder, image_file, max_width_inch=1.5):
    if image_file is None:
        return

    try:
        image_file.seek(0)
        image_bytes = image_file.read()
    except Exception as e:
        print("Error reading image file:", e)
        return

    with tempfile.NamedTemporaryFile(delete=False, suffix=".png") as tmp_img:
        tmp_img.write(image_bytes)
        tmp_img_path = tmp_img.name

    try:
        with Image.open(tmp_img_path) as img:
            width_px, height_px = img.size
    except Exception as e:
        print("Error opening image file:", e)
        os.remove(tmp_img_path)
        return

    aspect_ratio = height_px / width_px
    width = Inches(max_width_inch)
    height = Inches(max_width_inch * aspect_ratio)

    for para in paragraphs:
        if placeholder in para.text:
            para.clear()
            run = para.add_run()
            run.add_picture(tmp_img_path, width=width, height=height)

    try:
        os.remove(tmp_img_path)
    except Exception as e:
        print("Cannot delete temp file:", e)

def parse_thai_date_str(date_str):
    months_th = {
        "มกราคม": 1, "กุมภาพันธ์": 2, "มีนาคม": 3, "เมษายน": 4,
        "พฤษภาคม": 5, "มิถุนายน": 6, "กรกฎาคม": 7, "สิงหาคม": 8,
        "กันยายน": 9, "ตุลาคม": 10, "พฤศจิกายน": 11, "ธันวาคม": 12
    }
    try:
        parts = date_str.strip().split()
        day = int(parts[0])
        month = months_th.get(parts[1])
        year = int(parts[2]) - 543
        return datetime(year, month, day)
    except Exception:
        return None

def date_data(dt):
    return {
        "date_obj": dt,
        "day": dt.day,
        "month_th": format_date(dt, "MMMM", locale="th_TH"),
        "year_th": dt.year + 543,
        "day_ordinal": ordinal(dt.day),
        "month_en": dt.strftime("%B"),
        "year_en": dt.year,
    }

def get_image_size(image_file):
    image_file.seek(0)
    with Image.open(image_file) as img:
        return img.size

def safe_filename(s):
    return re.sub(r'[\\/*?:"<>|]', "", s).replace(" ", "_")

def convert_en_date_to_thai(date_str):
    try:
        date_obj = datetime.strptime(date_str.strip(), "%d %B %Y")
        thai_date = format_date(date_obj, format="d MMMM y", locale="th_TH")
        year_th = date_obj.year + 543
        thai_date_with_year = thai_date.replace(str(date_obj.year), str(year_th))
        return thai_date_with_year
    except Exception as e:
        return f"Invalid Format: {e}"
